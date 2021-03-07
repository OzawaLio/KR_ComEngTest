# 해당 코드를 허락없이 수정 및 배포 하는 것을 금지합니다
# made by. '금마한남'

import random
from docx import Document

document = Document()
documentAns = Document()

file1 = open("computer words.txt", 'rt', encoding='UTF-8')
size = int(input("size : "))

key = 'key'
value = 'value'
arr_key = []
arr_value = []
while 1:
    line = file1.readline()
    if line == '':
        break
    if line[0] == '#':
        arr_key.append(key)
        arr_value.append(value)
        key = ''
        value = ''

        key = line.replace('#','')

    if line[0] == ":":
        value = value + line.replace(':','') + ' '

file1.close()

# 초기값인 key, value 삭제
del arr_key[0]
del arr_value[0]


for i in range(0,size):
    num = random.randrange(0,len(arr_key))
    if num%2 == 0:
        document.add_paragraph().add_run(">> " + str(arr_key[num])+"\n\n\n")
    else:
        document.add_paragraph().add_run(">> " + str(arr_value[num]) + "\n\n\n")

    documentAns.add_paragraph().add_run("> "+str(arr_key[num]) +": "+str(arr_value[num])+"\n")

    # 이미 출력될 데이터를 배열에서 삭제
    del arr_key[num]
    del arr_value[num]

document.save("Computer Test.docx")
documentAns.save("Computer Test Answer.docx")
